"""
API路由 - 知识库管理
"""
import os
os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TRANSFORMERS_OFFLINE"] = "1"

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, func as sql_func
from typing import List, Optional, Dict, Any
import uuid
import json
import logging
from datetime import datetime

from app.core.database import get_db
from app.core.security import get_current_user
from app.core.config import settings
from app.models.user import User
from app.models.ai import KnowledgeBase, KnowledgeDocument, KnowledgeTag, KnowledgeDocumentTag, KnowledgeNote
from app.schemas.schemas import (
    KnowledgeBaseCreate, KnowledgeBaseUpdate, BaseResponse, DocumentUploadResponse,
    KnowledgeTagCreate, KnowledgeTagUpdate, DocumentTagRequest,
    KnowledgeNoteCreate, KnowledgeNoteUpdate
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/knowledge", tags=["知识库"])


# ============ 知识库CRUD ============

@router.get("/bases")
async def list_knowledge_bases(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取知识库列表"""
    query = select(KnowledgeBase).where(KnowledgeBase.is_active == True)
    result = await db.execute(query)
    bases = result.scalars().all()

    return [
        {
            "id": kb.id,
            "name": kb.name,
            "code": kb.code,
            "description": kb.description,
            "embedding_model": kb.embedding_model,
            "rerank_model": kb.rerank_model,
            "rerank_enabled": kb.rerank_enabled,
            "doc_count": kb.doc_count,
            "vector_count": kb.vector_count,
            "created_at": kb.created_at,
            "updated_at": kb.updated_at,
        }
        for kb in bases
    ]


@router.get("/bases/{kb_id}")
async def get_knowledge_base(
    kb_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取知识库详情"""
    result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    return {
        "id": kb.id,
        "name": kb.name,
        "code": kb.code,
        "description": kb.description,
        "embedding_model": kb.embedding_model,
        "rerank_model": kb.rerank_model,
        "rerank_enabled": kb.rerank_enabled,
        "config": kb.config,
        "doc_count": kb.doc_count,
        "vector_count": kb.vector_count,
        "is_active": kb.is_active,
        "created_at": kb.created_at,
        "updated_at": kb.updated_at,
    }


@router.post("/bases", response_model=BaseResponse)
async def create_knowledge_base(
    request: KnowledgeBaseCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """创建知识库"""
    code = request.code or f"kb_{uuid.uuid4().hex[:8]}"

    kb = KnowledgeBase(
        name=request.name,
        code=code,
        description=request.description,
        embedding_model=request.embedding_model,
        rerank_model=request.rerank_model,
        rerank_enabled=request.rerank_enabled or False,
        config=request.config or {},
        organization_id=current_user.organization_id,
        created_by=current_user.id
    )

    db.add(kb)
    await db.commit()
    await db.refresh(kb)

    return BaseResponse(message="知识库创建成功", data={
        "id": kb.id,
        "name": kb.name,
        "code": kb.code,
        "embedding_model": kb.embedding_model,
        "config": kb.config,
    })


@router.put("/bases/{kb_id}", response_model=BaseResponse)
async def update_knowledge_base(
    kb_id: int,
    request: KnowledgeBaseUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """更新知识库"""
    result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    if request.name is not None:
        kb.name = request.name
    if request.description is not None:
        kb.description = request.description
    if request.embedding_model is not None:
        kb.embedding_model = request.embedding_model
    if request.rerank_model is not None:
        kb.rerank_model = request.rerank_model
    if request.rerank_enabled is not None:
        kb.rerank_enabled = request.rerank_enabled
    if request.config is not None:
        # 合并配置
        kb.config = {**(kb.config or {}), **request.config}

    await db.commit()
    await db.refresh(kb)

    return BaseResponse(message="知识库更新成功", data={
        "id": kb.id,
        "name": kb.name,
        "embedding_model": kb.embedding_model,
        "config": kb.config,
    })


@router.delete("/bases/{kb_id}", response_model=BaseResponse)
async def delete_knowledge_base(
    kb_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """删除知识库（软删除）"""
    result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    kb.is_active = False
    await db.commit()

    return BaseResponse(message="知识库已删除")


# ============ 文档管理 ============

@router.get("/documents")
async def list_documents(
    kb_id: Optional[int] = Query(None, description="知识库ID过滤"),
    status: Optional[str] = Query(None, description="解析状态过滤"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取文档列表，支持kb_id过滤"""
    query = select(KnowledgeDocument).where(KnowledgeDocument.is_active == True)
    if kb_id is not None:
        query = query.where(KnowledgeDocument.knowledge_base_id == kb_id)
    if status is not None:
        query = query.where(KnowledgeDocument.parsing_status == status)

    query = query.order_by(KnowledgeDocument.created_at.desc())
    result = await db.execute(query)
    docs = result.scalars().all()

    return [
        {
            "id": doc.id,
            "knowledge_base_id": doc.knowledge_base_id,
            "title": doc.title,
            "content": doc.content[:500] + "..." if doc.content and len(doc.content) > 500 else doc.content,
            "file_name": doc.file_name,
            "file_size": doc.file_size,
            "file_type": doc.file_type,
            "chunk_count": doc.chunk_count,
            "parsing_status": doc.parsing_status,
            "parsing_error": doc.parsing_error,
            "keywords": doc.keywords if isinstance(doc.keywords, list) else [],
            "summary": doc.summary or "",
            "tags": [],
            "created_at": doc.created_at,
            "updated_at": doc.updated_at,
        }
        for doc in docs
    ]


@router.post("/upload/{kb_id}")
async def upload_document(
    kb_id: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """上传单个文档到知识库"""
    # 检查知识库
    result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    # 保存文件
    file_ext = os.path.splitext(file.filename)[1]
    file_id = uuid.uuid4().hex
    kb_upload_dir = os.path.join(settings.UPLOAD_DIR, f"kb_{kb_id}")
    os.makedirs(kb_upload_dir, exist_ok=True)
    file_path = os.path.join(kb_upload_dir, f"{file_id}{file_ext}")

    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)

    # 创建文档记录
    doc = KnowledgeDocument(
        knowledge_base_id=kb_id,
        title=file.filename,
        file_name=file.filename,
        file_path=file_path,
        file_size=len(content),
        file_type=file_ext[1:] if file_ext else "unknown",
        parsing_status="pending"
    )

    db.add(doc)
    kb.doc_count = (kb.doc_count or 0) + 1

    await db.commit()
    await db.refresh(doc)

    # 上传后立即解析（仅文本提取+关键词+摘要，不向量化）
    try:
        await _parse_document(doc.id, db)
    except Exception as e:
        logger.warning(f"自动解析失败: {e}")

    return DocumentUploadResponse(
        id=doc.id,
        title=doc.title,
        file_name=doc.file_name,
        file_size=doc.file_size,
        parsing_status=doc.parsing_status
    )


@router.post("/upload-batch/{kb_id}")
async def upload_document_batch(
    kb_id: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """批量上传文档到知识库（支持逐个上传）"""
    # 检查知识库
    result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    
    # 保存文件
    file_ext = os.path.splitext(file.filename)[1]
    file_id = uuid.uuid4().hex
    kb_upload_dir = os.path.join(settings.UPLOAD_DIR, f"kb_{kb_id}")
    os.makedirs(kb_upload_dir, exist_ok=True)
    file_path = os.path.join(kb_upload_dir, f"{file_id}{file_ext}")
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # 创建文档记录
    doc = KnowledgeDocument(
        knowledge_base_id=kb_id,
        title=file.filename,
        file_name=file.filename,
        file_path=file_path,
        file_size=len(content),
        file_type=file_ext[1:] if file_ext else "unknown",
        parsing_status="pending"
    )
    
    db.add(doc)
    kb.doc_count = (kb.doc_count or 0) + 1
    
    await db.commit()
    await db.refresh(doc)
    
    return {
        "id": doc.id,
        "title": doc.title,
        "file_name": doc.file_name,
        "file_size": doc.file_size,
        "parsing_status": doc.parsing_status
    }



@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """删除文档"""
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()

    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 删除文件
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    # 更新知识库统计
    kb_result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == doc.knowledge_base_id)
    )
    kb = kb_result.scalar_one_or_none()
    if kb and kb.doc_count and kb.doc_count > 0:
        kb.doc_count -= 1

    await db.delete(doc)
    await db.commit()

    return BaseResponse(message="文档已删除")


# ============ 文档解析 ============

@router.post("/parse/{doc_id}", response_model=BaseResponse)
async def parse_document(
    doc_id: int,
    vectorize: bool = Query(False, description="是否同时向量化"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """解析单个文档（默认仅文本解析，vectorize=true时同时向量化）"""
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=400, detail="文档文件不存在")

    try:
        await _parse_document(doc_id, db)
        if vectorize:
            await _vectorize_document(doc_id, db)
        return BaseResponse(message="文档解析完成", data={"doc_id": doc_id, "status": doc.parsing_status})
    except Exception as e:
        logger.error(f"解析文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"解析失败: {str(e)}")


@router.post("/parse-all/{kb_id}", response_model=BaseResponse)
async def parse_all_documents(
    kb_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """批量解析知识库所有未解析文档"""
    # 检查知识库
    kb_result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = kb_result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    # 查找所有待解析文档
    query = select(KnowledgeDocument).where(
        KnowledgeDocument.knowledge_base_id == kb_id,
        KnowledgeDocument.is_active == True,
        KnowledgeDocument.parsing_status.in_(["pending", "failed"])
    )
    result = await db.execute(query)
    docs = result.scalars().all()

    success_count = 0
    fail_count = 0
    errors = []

    for doc in docs:
        try:
            await _parse_document(doc.id, db)
            success_count += 1
        except Exception as e:
            fail_count += 1
            errors.append({"doc_id": doc.id, "title": doc.title, "error": str(e)})
            logger.error(f"批量解析文档 {doc.id} 失败: {e}")

    return BaseResponse(
        message=f"批量解析完成: 成功 {success_count} 个, 失败 {fail_count} 个",
        data={
            "total": len(docs),
            "success": success_count,
            "failed": fail_count,
            "errors": errors if errors else None,
        }
    )


@router.post("/vectorize/{doc_id}", response_model=BaseResponse)
async def vectorize_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """对已解析文档执行向量化"""
    doc_result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = doc_result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if doc.parsing_status == "completed":
        return BaseResponse(message="文档已向量化", data={"doc_id": doc_id})
    if doc.parsing_status not in ("parsed", "failed"):
        return BaseResponse(message="文档尚未完成文本解析，请先解析", data={"doc_id": doc_id, "status": doc.parsing_status})

    try:
        await _vectorize_document(doc_id, db)
        return BaseResponse(message="向量化完成", data={"doc_id": doc_id, "status": doc.parsing_status})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"向量化失败: {str(e)}")


@router.post("/vectorize-all/{kb_id}", response_model=BaseResponse)
async def vectorize_all_documents(
    kb_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """一键批量向量化知识库中所有已解析但未向量化的文档"""
    kb_result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
    )
    kb = kb_result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    query = select(KnowledgeDocument).where(
        KnowledgeDocument.knowledge_base_id == kb_id,
        KnowledgeDocument.is_active == True,
        KnowledgeDocument.parsing_status.in_(["parsed", "failed"])
    )
    result = await db.execute(query)
    docs = result.scalars().all()

    success_count = 0
    fail_count = 0
    errors = []

    for doc in docs:
        try:
            await _vectorize_document(doc.id, db)
            success_count += 1
        except Exception as e:
            fail_count += 1
            errors.append({"doc_id": doc.id, "title": doc.title, "error": str(e)})
            logger.error(f"向量化文档 {doc.id} 失败: {e}")

    return BaseResponse(
        message=f"批量向量化完成: 成功 {success_count} 个, 失败 {fail_count} 个",
        data={
            "total": len(docs),
            "success": success_count,
            "failed": fail_count,
            "errors": errors if errors else None,
        }
    )


# ============ 查询知识库 ============

@router.get("/documents/{doc_id}")
async def get_document_detail(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取文档完整内容"""
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    return {
        "id": doc.id,
        "knowledge_base_id": doc.knowledge_base_id,
        "title": doc.title,
        "content": doc.content,
        "file_name": doc.file_name,
        "file_size": doc.file_size,
        "file_type": doc.file_type,
        "chunk_count": doc.chunk_count,
        "parsing_status": doc.parsing_status,
        "parsing_error": doc.parsing_error,
        "created_at": doc.created_at,
        "updated_at": doc.updated_at,
    }


@router.post("/query")
async def query_knowledge(
    query: str,
    kb_id: Optional[int] = None,
    top_k: int = 5,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """查询知识库"""
    from app.core.ai_digital_base import rag_retriever

    collection_name = f"kb_{kb_id}" if kb_id else "global"

    # 获取 KB embedding 配置（支持 KB 指定的本地/API 模型）
    kb_config = None
    if kb_id:
        kb_result = await db.execute(
            select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
        )
        kb = kb_result.scalar_one_or_none()
        if kb:
            kb_config = _build_kb_embedding_config(kb)

    # 初始检索多取一些结果，便于后续重排
    search_top_k = top_k * 3 if kb_id else top_k
    results = await rag_retriever.search(
        collection_name=collection_name,
        query=query,
        top_k=search_top_k,
        kb_config=kb_config,
    )

    # Rerank：使用知识库配置的重排模型
    if kb_id and results:
        kb_result = await db.execute(
            select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
        )
        kb = kb_result.scalar_one_or_none()

        if kb and kb.rerank_enabled and kb.rerank_model:
            results = await _rerank_results(query, results, kb.rerank_model, top_k)
    
    # 截取到 top_k
    results = results[:top_k]

    return {"results": results}


@router.post("/app-search")
async def search_app_knowledge(
    app_id: int = Query(..., description="应用ID"),
    query: str = Query(..., description="查询文本"),
    top_k: int = Query(5, description="返回结果数"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    在应用绑定的知识库中搜索
    支持应用级自动索引的数据检索
    """
    from app.core.rag_autoindexer import get_rag_autoindexer
    
    autoindexer = get_rag_autoindexer()
    
    results = await autoindexer.search_app_knowledge(
        app_id=app_id,
        query=query,
        top_k=top_k,
        db=db,
    )
    
    return {
        "app_id": app_id,
        "query": query,
        "results": results,
        "total": len(results)
    }


# ============ 内部解析函数 ============

async def _parse_document(doc_id: int, db: AsyncSession):
    """
    快速解析文档：提取文字 → jieba分词/关键词/摘要
    上传后自动调用，不包含向量化（向量化由用户手动触发）
    """
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        return

    doc.parsing_status = "processing"
    await db.commit()

    try:
        # 1. 提取文字
        text = _extract_text(doc.file_path, doc.file_type)

        if not text or not text.strip():
            doc.parsing_status = "failed"
            doc.parsing_error = "未能从文档中提取到文字内容"
            await db.commit()
            return

        # 2. jieba分词、关键词、摘要
        keywords = []
        summary = ""
        try:
            from app.core.ai_digital_base.local_services import text_parser_service

            kw_result = text_parser_service.extract_keywords(text, top_k=10, method="tfidf")
            if kw_result.get("success"):
                keywords = [k["word"] for k in kw_result.get("keywords", [])]

            summary_result = text_parser_service.extract_summary(text, max_length=200)
            if summary_result.get("success"):
                summary = summary_result.get("summary", "")
        except Exception as kw_err:
            logger.warning(f"关键词/摘要提取失败(不影响解析): {kw_err}")

        # 3. 更新文档内容和状态
        doc.content = text
        doc.keywords = keywords
        doc.summary = summary
        doc.parsing_status = "parsed"  # parsed = 已解析文本，未向量化
        doc.chunk_count = len(_split_text(text, chunk_size=500, overlap=50))

        await db.commit()

    except Exception as e:
        doc.parsing_status = "failed"
        doc.parsing_error = str(e)
        await db.commit()
        logger.error(f"文档解析异常 doc_id={doc_id}: {e}")
        raise


async def _vectorize_document(doc_id: int, db: AsyncSession):
    """
    向量化文档：文本分块 → embedding → 向量存储
    由用户手动触发，耗时较长
    根据知识库配置选择正确的 embedding 模型（本地或API）
    """
    result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        return False

    if not doc.content or not doc.content.strip():
        logger.warning(f"文档 {doc_id} 无内容，跳过向量化")
        return False

    # 获取知识库配置（包含 embedding 模型配置）
    kb_result = await db.execute(
        select(KnowledgeBase).where(KnowledgeBase.id == doc.knowledge_base_id)
    )
    kb = kb_result.scalar_one_or_none()
    if not kb:
        logger.error(f"知识库 {doc.knowledge_base_id} 不存在，跳过向量化")
        return False

    # 构建 KB 的 embedding 配置
    kb_config = _build_kb_embedding_config(kb)
    logger.info(f"文档 {doc_id} 使用 embedding 模型: {kb_config.get('embedding_model')}, "
                f"类型: {kb_config.get('embedding_model_type')}")

    doc.parsing_status = "vectorizing"
    await db.commit()

    try:
        chunks = _split_text(doc.content, chunk_size=500, overlap=50)
        doc.chunk_count = len(chunks)

        from app.core.ai_digital_base import rag_retriever

        collection_name = f"kb_{doc.knowledge_base_id}"
        vector_count = 0

        keywords = doc.keywords if isinstance(doc.keywords, list) else []

        for i, chunk in enumerate(chunks):
            chunk_id = f"doc_{doc.id}_chunk_{i}"
            chunk_metadata = {
                "doc_id": doc.id,
                "doc_title": doc.title,
                "chunk_index": i,
                "keywords": keywords[:5] if keywords else [],
                "summary": doc.summary or "",
            }
            success = await rag_retriever.add_document(
                collection_name=collection_name,
                doc_id=chunk_id,
                text=chunk,
                metadata=chunk_metadata,
                kb_config=kb_config,
            )
            if success:
                vector_count += 1

        # 更新知识库向量计数
        kb.vector_count = (kb.vector_count or 0) + vector_count

        doc.parsing_status = "completed"  # completed = 已向量化
        await db.commit()
        logger.info(f"文档 {doc_id} 向量化完成（共 {vector_count}/{len(chunks)} 个向量），模型: {kb_config.get('embedding_model')}")
        return True

    except Exception as e:
        doc.parsing_status = "parsed"  # 向量化失败但文本解析保留
        doc.parsing_error = f"向量化失败: {str(e)}"
        await db.commit()
        logger.error(f"文档向量化异常 doc_id={doc_id}: {e}")
        raise


def _build_kb_embedding_config(kb) -> Dict[str, Any]:
    """
    根据知识库配置构建 embedding 配置字典
    优先使用 KB.embedding_model 字段，再合并 config 中的详细配置
    """
    # embedding_model 字段是基础模型名称
    embedding_model = kb.embedding_model or "text-embedding-v2"
    
    # config 中有更详细的配置（模型类型、路径等）
    kb_config = kb.config or {}
    
    config = {
        "embedding_model": embedding_model,
        # 优先读 config 中的类型配置，否则根据模型名推断
        "embedding_model_type": kb_config.get("embedding_model_type") or _infer_embedding_type(embedding_model),
        "embedding_model_path": kb_config.get("embedding_model_path"),
        "embedding_api_key": kb_config.get("embedding_api_key"),
        "embedding_api_base": kb_config.get("embedding_api_base"),
    }
    
    # 如果 config 中有 API 密钥覆盖全局设置
    if kb_config.get("embedding_api_key"):
        config["embedding_api_key"] = kb_config["embedding_api_key"]
    if kb_config.get("embedding_api_base"):
        config["embedding_api_base"] = kb_config["embedding_api_base"]
    
    return config


def _infer_embedding_type(model_name: str) -> str:
    """
    根据模型名称推断 embedding 类型
    """
    if not model_name:
        return "api"
    
    model_lower = model_name.lower()
    
    # 明确是本地路径
    if model_name.startswith(("E:\\", "C:\\", "/", "D:\\")):
        return "local"
    
    # sentence-transformers 内置轻量模型
    local_models = [
        "all-MiniLM-L6-v2",
        "paraphrase-multilingual-MiniLM-L12-v2",
        "paraphrase-multilingual-mpnet-base-v2",
        "shibing624/text2vec-base-chinese",
        "DMetaSoul/sbert-chinese-qmc-domain-v1",
        "moka-ai/m3e-small",
        "moka-ai/m3e-base",
        "moka-ai/m3e-large",
        "BAAI/bge-small-zh-v1.5",
        "BAAI/bge-base-zh-v1.5",
        "BAAI/bge-large-zh-v1.5",
        "sentence-transformers/",
    ]
    
    for lm in local_models:
        if lm in model_lower or model_lower in lm:
            return "local"
    
    # API 模型通常是 dashscope/openai 等服务
    return "api"


def _extract_text(file_path: str, file_type: str) -> str:
    """根据文件类型提取文字"""
    file_type = (file_type or "").lower()

    # 图片 -> OCR
    if file_type in ("jpg", "jpeg", "png", "bmp", "tiff", "gif", "webp"):
        return _extract_text_from_image(file_path)

    # Word文档
    if file_type in ("docx", "doc"):
        return _extract_text_from_docx(file_path)

    # PDF
    if file_type == "pdf":
        return _extract_text_from_pdf(file_path)

    # Excel
    if file_type in ("xlsx", "xls"):
        return _extract_text_from_excel(file_path)

    # 纯文本 / Markdown
    if file_type in ("txt", "md", "markdown", "csv", "json", "xml", "html", "log"):
        return _extract_text_from_txt(file_path)

    # 未知类型尝试按文本读取
    try:
        return _extract_text_from_txt(file_path)
    except Exception:
        return ""


def _extract_text_from_image(file_path: str) -> str:
    """使用OCR提取图片文字"""
    from app.core.ai_digital_base.local_services import ocr_service, TESSERACT_AVAILABLE
    
    # 检查 OCR 是否可用
    if not TESSERACT_AVAILABLE:
        logger.warning("OCR提取失败: pytesseract 未安装")
        return ""
    
    if not ocr_service.is_configured():
        logger.warning("OCR提取失败: Tesseract 未配置或不可用")
        return ""
    
    with open(file_path, "rb") as f:
        image_data = f.read()
    
    result = ocr_service.extract_text(image_data)
    if result.get("success"):
        return result.get("text", "")
    else:
        logger.warning(f"OCR提取失败: {result.get('error', '未知错误')}")
        return ""


def _extract_text_from_docx(file_path: str) -> str:
    """从Word文档提取文字"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # 也提取表格内容
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))
        return "\n".join(paragraphs + tables_text)
    except ImportError:
        logger.warning("python-docx 未安装，无法解析Word文档")
        return ""
    except Exception as e:
        logger.error(f"Word文档解析失败: {e}")
        return ""


def _extract_text_from_pdf(file_path: str) -> str:
    """从PDF提取文字"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except ImportError:
        logger.warning("pypdf 未安装，无法解析PDF文档")
        return ""
    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        return ""


def _extract_text_from_excel(file_path: str) -> str:
    """从Excel提取数据"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        rows_text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    rows_text.append(" | ".join(row_text))
        wb.close()
        return "\n".join(rows_text)
    except ImportError:
        logger.warning("openpyxl 未安装，无法解析Excel文档")
        return ""
    except Exception as e:
        logger.error(f"Excel解析失败: {e}")
        return ""


def _extract_text_from_txt(file_path: str) -> str:
    """读取纯文本文件"""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""


def _split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    文本分块
    按段落/句子分割，尽量保持语义完整
    """
    if not text:
        return []

    # 先按段落分割
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    chunks = []
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 <= chunk_size:
            current_chunk = f"{current_chunk}\n{para}" if current_chunk else para
        else:
            if current_chunk:
                chunks.append(current_chunk)
            # 如果单段落超长，按句子再分
            if len(para) > chunk_size:
                sentences = _split_by_sentences(para)
                sent_chunk = ""
                for sent in sentences:
                    if len(sent_chunk) + len(sent) + 1 <= chunk_size:
                        sent_chunk = f"{sent_chunk}{sent}" if sent_chunk else sent
                    else:
                        if sent_chunk:
                            chunks.append(sent_chunk)
                        sent_chunk = sent
                if sent_chunk:
                    current_chunk = sent_chunk
                else:
                    current_chunk = ""
            else:
                current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    # 如果没有成功分块，强制按字符分
    if not chunks and text:
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if chunk.strip():
                chunks.append(chunk)

    return chunks


def _split_by_sentences(text: str) -> List[str]:
    """按句子分割文本"""
    import re
    sentences = re.split(r'([。！？；\.\!\?;])', text)
    result = []
    for i in range(0, len(sentences) - 1, 2):
        sent = sentences[i] + (sentences[i + 1] if i + 1 < len(sentences) else "")
        if sent.strip():
            result.append(sent)
    if len(sentences) % 2 == 1 and sentences[-1].strip():
        result.append(sentences[-1])
    return result


# CrossEncoder 模型缓存（避免每次查询重新加载）
_cross_encoder_cache: Dict[str, Any] = {}

# 本地模型路径配置（支持多个位置）
LOCAL_RERANKER_MODEL_PATHS = [
    os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
        "bge-reranker-v2-m3"
    ),
    r"E:\models\bge-reranker-v2-m3",
]


def _get_reranker_model_path(model_name: str) -> str:
    """
    获取 reranker 模型的本地路径
    优先使用本地已存在的模型，否则从魔塔社区下载
    """
    from app.core.ai_digital_base.modelscope_utils import ensure_model_downloaded
    from app.core.config import settings
    
    # 如果 model_name 本身就是一个存在的本地路径
    if os.path.exists(model_name):
        has_weights = (
            os.path.exists(os.path.join(model_name, "model.safetensors")) or
            os.path.exists(os.path.join(model_name, "pytorch_model.bin"))
        )
        if has_weights:
            logger.info(f"使用本地 reranker 模型: {model_name}")
            return model_name
    
    # 检查是否是 bge-reranker-v2-m3 模型
    if "bge-reranker-v2-m3" in model_name:
        for local_path in LOCAL_RERANKER_MODEL_PATHS:
            if os.path.exists(local_path):
                has_weights = (
                    os.path.exists(os.path.join(local_path, "model.safetensors")) or
                    os.path.exists(os.path.join(local_path, "pytorch_model.bin"))
                )
                if has_weights:
                    logger.info(f"使用本地 reranker 模型: {local_path}")
                    return local_path
    
    # 尝试从魔塔社区下载模型
    cache_dir = os.path.join(settings.PROJECT_ROOT, "models", "cache")
    logger.info(f"尝试从魔塔社区获取 rerank 模型: {model_name}")
    
    downloaded_path = ensure_model_downloaded(model_name, cache_dir=cache_dir)
    
    if downloaded_path and os.path.exists(downloaded_path):
        # 验证模型文件是否完整
        has_weights = (
            os.path.exists(os.path.join(downloaded_path, "model.safetensors")) or
            os.path.exists(os.path.join(downloaded_path, "pytorch_model.bin"))
        )
        if has_weights:
            logger.info(f"使用魔塔社区 rerank 模型: {downloaded_path}")
            return downloaded_path
    
    raise FileNotFoundError(
        f"无法获取 reranker 模型: {model_name}\n"
        f"本地路径不存在且从魔塔社区下载失败。\n"
        f"请检查网络连接或手动下载模型到以下路径之一:\n" +
        "\n".join(f"  - {p}" for p in LOCAL_RERANKER_MODEL_PATHS)
    )


async def _rerank_results(query: str, results: List[Dict], rerank_model: str, top_k: int) -> List[Dict]:
    """
    使用AI模型对检索结果进行重排
    支持两种模式：
    1. rerank_model 为已知 rerank 模型名（如 BAAI/bge-reranker-v2-m3）→ 使用 sentence-transformers CrossEncoder（必须本地存在，禁止下载）
    2. rerank_model 为系统配置的AI模型ID → 调用 LLM 对结果重新打分排序
    """
    if not results:
        return results

    # 尝试使用 CrossEncoder 本地 rerank（禁止自动下载）
    try:
        from sentence_transformers import CrossEncoder

        # 获取模型路径（必须本地存在）
        model_path = _get_reranker_model_path(rerank_model)

        if model_path not in _cross_encoder_cache:
            logger.info(f"加载 CrossEncoder rerank 模型: {model_path}")
            _cross_encoder_cache[model_path] = CrossEncoder(
                model_path,
                max_length=512,
                device="cpu"
            )
            logger.info("CrossEncoder 模型加载成功")

        ce = _cross_encoder_cache[model_path]
        pairs = [(query, r.get("text", r.get("content", ""))) for r in results]

        logger.debug(f"Reranking {len(pairs)} documents...")
        scores = ce.predict(pairs)

        for i, r in enumerate(results):
            r["rerank_score"] = float(scores[i])

        results.sort(key=lambda x: x.get("rerank_score", 0), reverse=True)
        logger.info(f"CrossEncoder rerank 完成，模型: {model_path}")
        return results[:top_k]

    except FileNotFoundError as e:
        logger.warning(f"CrossEncoder 本地模型未找到，降级到 LLM 重排: {e}")
    except ImportError:
        logger.warning("sentence-transformers 未安装，CrossEncoder 不可用，尝试 LLM rerank")
    except Exception as e:
        logger.warning(f"CrossEncoder rerank 失败，降级到 LLM 重排: {e}")
    
    # 使用系统配置的AI模型进行LLM重排
    try:
        from app.core.ai_digital_base.gateway import ai_gateway
        await ai_gateway.load_config_from_db()
        
        # 构建重排 prompt
        docs_text = ""
        for i, r in enumerate(results[:10]):  # 最多10个
            docs_text += f"\n[{i+1}] {r.get('text', r.get('content', ''))[:300]}"
        
        prompt = f"""请对以下检索结果按与查询的相关性从高到低排序，只返回排序后的编号列表（如：3,1,5,2,4），不要其他内容。

查询: {query}

检索结果:{docs_text}

排序结果:"""

        resp = await ai_gateway.chat(
            messages=[{"role": "user", "content": prompt}],
            model=rerank_model,
            max_tokens=100,
            temperature=0.1
        )
        
        if "error" not in resp and resp.get("content"):
            # 解析排序结果
            import re
            nums = re.findall(r'\d+', resp["content"])
            if nums:
                ranked_indices = [int(n) - 1 for n in nums if 0 < int(n) <= len(results)]
                # 添加未在排序中的结果
                remaining = [i for i in range(len(results)) if i not in ranked_indices]
                all_indices = ranked_indices + remaining
                results = [results[i] for i in all_indices if i < len(results)]
                return results[:top_k]
    except Exception as e:
        logger.warning(f"LLM rerank 失败: {e}")
    
    # 降级：返回原始排序
    return results[:top_k]


# ============ 高级检索 ============

def _compute_bm25_score(query_terms: List[str], doc_text: str, avg_doc_len: float = 100) -> float:
    """简化的BM25评分"""
    if not query_terms or not doc_text:
        return 0.0
    k1, b = 1.5, 0.75
    doc_len = len(doc_text)
    score = 0.0
    doc_lower = doc_text.lower()
    for term in query_terms:
        tf = doc_lower.count(term.lower())
        if tf > 0:
            idf = 1.0  # 简化IDF
            score += idf * (tf * (k1 + 1)) / (tf + k1 * (1 - b + b * doc_len / avg_doc_len))
    return score


def _extract_tags_from_json(tags_json: Any) -> List[str]:
    """从JSON字段提取标签名列表"""
    if not tags_json:
        return []
    if isinstance(tags_json, str):
        try:
            tags_json = json.loads(tags_json)
        except:
            return [tags_json] if tags_json else []
    if isinstance(tags_json, list):
        return [t.get('name', t) if isinstance(t, dict) else str(t) for t in tags_json]
    return []


async def _do_vector_search(query: str, kb_id: Optional[int], top_k: int, kb_config: Dict[str, Any] = None) -> List[Dict[str, Any]]:
    """独立的向量搜索函数，支持 KB 指定的 embedding 模型"""
    try:
        from app.core.ai_digital_base import rag_retriever
        collection_name = f"kb_{kb_id}" if kb_id else "global"
        return await rag_retriever.search(
            collection_name=collection_name,
            query=query,
            top_k=top_k,
            kb_config=kb_config,
        )
    except Exception as e:
        logger.warning(f"向量检索异常: {e}")
        return []


@router.get("/search")
async def advanced_search(
    q: str = Query(..., description="搜索关键词"),
    type: str = Query("hybrid", description="检索类型: fulltext/keyword/vector/hybrid"),
    kb_id: Optional[int] = Query(None, description="知识库ID"),
    tag: Optional[str] = Query(None, description="标签过滤"),
    top_k: int = Query(10, description="返回数量"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """高级检索：支持全文/关键词/向量/混合四种模式，使用RRF融合"""
    import jieba
    import asyncio
    from sqlalchemy import or_, func
    
    results = {}  # 用dict存储，key=doc_id，便于RRF融合
    query_terms = [w for w in jieba.cut(q) if len(w) > 1]
    
    # ========== 1. 全文检索（BM25评分）==========
    if type in ("fulltext", "hybrid"):
        query_obj = select(KnowledgeDocument).where(KnowledgeDocument.is_active == True)
        if kb_id:
            query_obj = query_obj.where(KnowledgeDocument.knowledge_base_id == kb_id)
        if tag:
            # 标签过滤：检查tags JSON数组是否包含该标签
            query_obj = query_obj.where(
                or_(
                    KnowledgeDocument.tags.astext.like(f'%"name": "{tag}"%'),
                    KnowledgeDocument.tags.astext.like(f'%"{tag}"%')
                )
            )
        # 全文匹配：标题 OR 内容 OR 摘要
        query_obj = query_obj.where(
            or_(
                KnowledgeDocument.title.ilike(f"%{q}%"),
                KnowledgeDocument.content.ilike(f"%{q}%"),
                KnowledgeDocument.summary.ilike(f"%{q}%")
            )
        ).limit(top_k * 3)
        result = await db.execute(query_obj)
        docs = result.scalars().all()
        for rank, doc in enumerate(docs):
            if doc.id not in results:
                results[doc.id] = {
                    "id": doc.id, "title": doc.title,
                    "kb_id": doc.knowledge_base_id,
                    "keywords": doc.keywords, "tags": doc.tags,
                    "summary": (doc.summary or "")[:200],
                    "content": (doc.content or "")[:500],
                    "created_at": str(doc.created_at) if doc.created_at else None,
                    "scores": {}, "rrf_score": 0.0
                }
            # BM25评分
            bm25 = _compute_bm25_score(query_terms, (doc.title or "") + " " + (doc.content or ""))
            results[doc.id]["scores"]["fulltext"] = bm25
            results[doc.id]["scores"]["fulltext_rank"] = rank + 1
    
    # ========== 2. 关键词检索（OR逻辑，jieba分词）==========
    if type in ("keyword", "hybrid") and query_terms:
        # 构建OR条件
        or_conditions = []
        for kw in query_terms[:8]:  # 取前8个关键词
            or_conditions.append(KnowledgeDocument.title.ilike(f"%{kw}%"))
            or_conditions.append(KnowledgeDocument.summary.ilike(f"%{kw}%"))
            or_conditions.append(KnowledgeDocument.content.ilike(f"%{kw}%"))
        
        if or_conditions:
            query_obj = select(KnowledgeDocument).where(KnowledgeDocument.is_active == True)
            if kb_id:
                query_obj = query_obj.where(KnowledgeDocument.knowledge_base_id == kb_id)
            query_obj = query_obj.where(or_(*or_conditions)).limit(top_k * 3)
            result = await db.execute(query_obj)
            docs = result.scalars().all()
            for rank, doc in enumerate(docs):
                if doc.id not in results:
                    results[doc.id] = {
                        "id": doc.id, "title": doc.title,
                        "kb_id": doc.knowledge_base_id,
                        "keywords": doc.keywords, "tags": doc.tags,
                        "summary": (doc.summary or "")[:200],
                        "content": (doc.content or "")[:500],
                        "created_at": str(doc.created_at) if doc.created_at else None,
                        "scores": {}, "rrf_score": 0.0
                    }
                # 计算匹配的关键词数量作为分数
                matched = sum(1 for kw in query_terms if kw.lower() in (doc.title or "").lower() or kw.lower() in (doc.content or "").lower())
                results[doc.id]["scores"]["keyword"] = matched / len(query_terms) if query_terms else 0
                results[doc.id]["scores"]["keyword_rank"] = rank + 1
    
    # ========== 3. 向量检索 ==========
    if type in ("vector", "hybrid"):
        # 获取 KB embedding 配置（支持 KB 指定的本地/API 模型）
        kb_config = None
        if kb_id:
            kb_result = await db.execute(
                select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
            )
            kb = kb_result.scalar_one_or_none()
            if kb:
                kb_config = _build_kb_embedding_config(kb)
        
        try:
            vec_results = await asyncio.wait_for(
                _do_vector_search(q, kb_id, top_k * 2, kb_config=kb_config),
                timeout=8.0
            )
            if vec_results:
                for rank, vr in enumerate(vec_results):
                    doc_id = vr.get("metadata", {}).get("doc_id")
                    if doc_id:
                        if doc_id not in results:
                            results[doc_id] = {
                                "id": doc_id,
                                "title": vr.get("metadata", {}).get("doc_title", "未知文档"),
                                "kb_id": kb_id,
                                "keywords": vr.get("metadata", {}).get("keywords", []),
                                "tags": [],
                                "summary": "",
                                "content": vr.get("text", "")[:500],
                                "created_at": None,
                                "scores": {}, "rrf_score": 0.0
                            }
                        results[doc_id]["scores"]["vector"] = vr.get("score", 0.5)
                        results[doc_id]["scores"]["vector_rank"] = rank + 1
        except asyncio.TimeoutError:
            logger.warning("向量检索超时，跳过")
        except Exception as e:
            logger.warning(f"向量检索失败: {e}")
    
    # ========== 4. RRF融合评分 ==========
    k = 60  # RRF常数
    for doc_id, doc in results.items():
        rrf = 0.0
        for score_type in ["fulltext", "keyword", "vector"]:
            rank_key = f"{score_type}_rank"
            if rank_key in doc["scores"]:
                rrf += 1.0 / (k + doc["scores"][rank_key])
        doc["rrf_score"] = rrf
        # 综合分数（RRF + 原始分数加权）
        doc["score"] = rrf * 0.6 + max(doc["scores"].get("fulltext", 0), doc["scores"].get("keyword", 0), doc["scores"].get("vector", 0)) * 0.4
    
    # ========== 5. Rerank重排（可选）==========
    sorted_results = sorted(results.values(), key=lambda x: x["score"], reverse=True)
    if kb_id and len(sorted_results) > top_k:
        kb_result = await db.execute(
            select(KnowledgeBase).where(KnowledgeBase.id == kb_id, KnowledgeBase.is_active == True)
        )
        kb = kb_result.scalar_one_or_none()
        if kb and kb.rerank_enabled and kb.rerank_model:
            try:
                sorted_results = await _rerank_results(q, sorted_results, kb.rerank_model, top_k * 2)
            except Exception as e:
                logger.warning(f"Rerank失败: {e}")
    
    # ========== 6. 返回结果 ==========
    sorted_results = sorted(sorted_results, key=lambda x: x.get("rerank_score", x.get("score", 0)), reverse=True)
    
    # 清理内部字段
    for r in sorted_results:
        r.pop("scores", None)
        r.pop("rrf_score", None)
        r.pop("content", None)  # 不返回完整内容
    
    return {"results": sorted_results[:top_k], "total": len(sorted_results)}


# ============ 标签管理 ============

@router.get("/tags")
async def list_tags(
    kb_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取标签列表"""
    query = select(KnowledgeTag)
    if kb_id:
        query = query.where((KnowledgeTag.kb_id == kb_id) | (KnowledgeTag.kb_id == None))
    result = await db.execute(query.order_by(KnowledgeTag.name))
    tags = result.scalars().all()
    return [
        {"id": t.id, "name": t.name, "color": t.color,
         "description": t.description, "kb_id": t.kb_id,
         "created_at": str(t.created_at) if t.created_at else None}
        for t in tags
    ]


@router.post("/tags", response_model=BaseResponse)
async def create_tag(
    request: KnowledgeTagCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """创建标签"""
    tag = KnowledgeTag(
        name=request.name, color=request.color or "#1890ff",
        description=request.description, kb_id=request.kb_id,
        created_by=current_user.id
    )
    db.add(tag)
    await db.commit()
    await db.refresh(tag)
    return BaseResponse(message="标签创建成功", data={"id": tag.id, "name": tag.name})


@router.delete("/tags/{tag_id}", response_model=BaseResponse)
async def delete_tag(
    tag_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """删除标签"""
    result = await db.execute(select(KnowledgeTag).where(KnowledgeTag.id == tag_id))
    tag = result.scalar_one_or_none()
    if not tag:
        raise HTTPException(status_code=404, detail="标签不存在")
    # 删除关联
    await db.execute(
        update(KnowledgeDocumentTag).where(KnowledgeDocumentTag.tag_id == tag_id).values(tag_id=None)
    )
    await db.delete(tag)
    await db.commit()
    return BaseResponse(message="标签已删除")


@router.post("/documents/{doc_id}/tags", response_model=BaseResponse)
async def add_document_tag(
    doc_id: int,
    request: DocumentTagRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """给文档添加标签"""
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    tag_result = await db.execute(select(KnowledgeTag).where(KnowledgeTag.id == request.tag_id))
    tag = tag_result.scalar_one_or_none()
    if not tag:
        raise HTTPException(status_code=404, detail="标签不存在")

    # 检查是否已关联
    exist = await db.execute(
        select(KnowledgeDocumentTag).where(
            KnowledgeDocumentTag.document_id == doc_id,
            KnowledgeDocumentTag.tag_id == request.tag_id
        )
    )
    if exist.scalar_one_or_none():
        return BaseResponse(message="标签已存在")

    rel = KnowledgeDocumentTag(document_id=doc_id, tag_id=request.tag_id)
    db.add(rel)

    # 同步更新文档tags JSON字段
    doc_tags = doc.tags or []
    if tag.name not in doc_tags:
        doc_tags.append(tag.name)
        doc.tags = doc_tags

    await db.commit()
    return BaseResponse(message="标签添加成功")


@router.delete("/documents/{doc_id}/tags/{tag_id}", response_model=BaseResponse)
async def remove_document_tag(
    doc_id: int,
    tag_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """移除文档标签"""
    result = await db.execute(
        select(KnowledgeDocumentTag).where(
            KnowledgeDocumentTag.document_id == doc_id,
            KnowledgeDocumentTag.tag_id == tag_id
        )
    )
    rel = result.scalar_one_or_none()
    if rel:
        await db.delete(rel)

    # 同步更新文档tags JSON
    doc_result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = doc_result.scalar_one_or_none()
    if doc and doc.tags:
        tag_result = await db.execute(select(KnowledgeTag).where(KnowledgeTag.id == tag_id))
        tag = tag_result.scalar_one_or_none()
        if tag and tag.name in doc.tags:
            doc.tags = [t for t in doc.tags if t != tag.name]

    await db.commit()
    return BaseResponse(message="标签已移除")


# ============ 笔记管理 ============

@router.get("/notes")
async def list_notes(
    kb_id: Optional[int] = Query(None),
    is_daily: Optional[bool] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取笔记列表"""
    query = select(KnowledgeNote).where(KnowledgeNote.is_active == True)
    if kb_id:
        query = query.where(KnowledgeNote.knowledge_base_id == kb_id)
    if is_daily is not None:
        query = query.where(KnowledgeNote.is_daily == is_daily)
    query = query.order_by(KnowledgeNote.updated_at.desc())
    result = await db.execute(query)
    notes = result.scalars().all()
    return [
        {"id": n.id, "title": n.title, "content": (n.content or "")[:500],
         "tags": n.tags, "is_daily": n.is_daily, "note_date": n.note_date,
         "knowledge_base_id": n.knowledge_base_id,
         "created_at": str(n.created_at) if n.created_at else None,
         "updated_at": str(n.updated_at) if n.updated_at else None}
        for n in notes
    ]


@router.post("/notes", response_model=BaseResponse)
async def create_note(
    request: KnowledgeNoteCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """创建笔记"""
    note = KnowledgeNote(
        title=request.title, content=request.content or "",
        tags=request.tags or [], is_daily=request.is_daily or False,
        note_date=request.note_date, knowledge_base_id=request.knowledge_base_id,
        organization_id=current_user.organization_id, created_by=current_user.id
    )
    db.add(note)
    await db.commit()
    await db.refresh(note)
    return BaseResponse(message="笔记创建成功", data={"id": note.id})


@router.put("/notes/{note_id}", response_model=BaseResponse)
async def update_note(
    note_id: int,
    request: KnowledgeNoteUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """更新笔记"""
    result = await db.execute(select(KnowledgeNote).where(KnowledgeNote.id == note_id))
    note = result.scalar_one_or_none()
    if not note:
        raise HTTPException(status_code=404, detail="笔记不存在")
    if request.title is not None:
        note.title = request.title
    if request.content is not None:
        note.content = request.content
    if request.tags is not None:
        note.tags = request.tags
    if request.is_daily is not None:
        note.is_daily = request.is_daily
    await db.commit()
    return BaseResponse(message="笔记更新成功")


@router.delete("/notes/{note_id}", response_model=BaseResponse)
async def delete_note(
    note_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """删除笔记"""
    result = await db.execute(select(KnowledgeNote).where(KnowledgeNote.id == note_id))
    note = result.scalar_one_or_none()
    if not note:
        raise HTTPException(status_code=404, detail="笔记不存在")
    note.is_active = False
    await db.commit()
    return BaseResponse(message="笔记已删除")


@router.get("/notes/{note_id}")
async def get_note_detail(
    note_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取笔记详情"""
    result = await db.execute(select(KnowledgeNote).where(KnowledgeNote.id == note_id))
    note = result.scalar_one_or_none()
    if not note:
        raise HTTPException(status_code=404, detail="笔记不存在")
    return {
        "id": note.id, "title": note.title, "content": note.content,
        "tags": note.tags, "is_daily": note.is_daily, "note_date": note.note_date,
        "knowledge_base_id": note.knowledge_base_id,
        "created_at": str(note.created_at) if note.created_at else None,
        "updated_at": str(note.updated_at) if note.updated_at else None,
    }


# ============ 知识图谱 ============

@router.get("/graph")
async def knowledge_graph(
    kb_id: Optional[int] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """知识图谱数据"""
    nodes = []
    edges = []

    # 知识库节点
    kb_query = select(KnowledgeBase).where(KnowledgeBase.is_active == True)
    if kb_id:
        kb_query = kb_query.where(KnowledgeBase.id == kb_id)
    result = await db.execute(kb_query)
    kbs = result.scalars().all()
    for kb in kbs:
        nodes.append({"id": f"kb_{kb.id}", "name": kb.name, "category": "知识库", "symbolSize": 40})

    # 文档节点（取前100）
    doc_query = select(KnowledgeDocument).where(KnowledgeDocument.is_active == True)
    if kb_id:
        doc_query = doc_query.where(KnowledgeDocument.knowledge_base_id == kb_id)
    doc_query = doc_query.limit(100)
    result = await db.execute(doc_query)
    docs = result.scalars().all()
    doc_ids = set()
    for doc in docs:
        doc_ids.add(doc.id)
        nodes.append({"id": f"doc_{doc.id}", "name": doc.title[:20], "category": "文档", "symbolSize": 20})
        edges.append({"source": f"kb_{doc.knowledge_base_id}", "target": f"doc_{doc.id}"})

    # 标签节点
    tag_query = select(KnowledgeTag)
    result = await db.execute(tag_query)
    tags = result.scalars().all()
    for tag in tags:
        nodes.append({"id": f"tag_{tag.id}", "name": tag.name, "category": "标签", "symbolSize": 15})

    # 文档-标签关联
    if doc_ids:
        rel_query = select(KnowledgeDocumentTag).where(KnowledgeDocumentTag.document_id.in_(doc_ids))
        result = await db.execute(rel_query)
        rels = result.scalars().all()
        for rel in rels:
            edges.append({"source": f"doc_{rel.document_id}", "target": f"tag_{rel.tag_id}"})

    return {"nodes": nodes, "edges": edges}
