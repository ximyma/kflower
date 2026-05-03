# -*- coding: utf-8 -*-
"""
智能导入服务 - 支持 Excel/Word/图片/CSV 文件解析生成表单字段
全新版本：支持智能表头检测、用户自选表头行
"""
import re
import json
import io
import os
from typing import List, Dict, Any, Tuple, Optional

# ============ 依赖检查 ============
EXCEL_AVAILABLE = False
WORD_AVAILABLE = False
JIEBA_AVAILABLE = False
TESSERACT_AVAILABLE = False
XLRD_AVAILABLE = False
CV2_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    pass

try:
    import docx
    import docx2txt
    WORD_AVAILABLE = True
except ImportError:
    pass

try:
    import jieba
    import jieba.analyse
    JIEBA_AVAILABLE = True
except ImportError:
    pass

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    pass

try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    pass

try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    pass


# ============ 文件类型检测 ============
def detect_file_type(file_bytes: bytes, filename: str = "") -> str:
    """
    根据文件头（Magic Bytes）和扩展名检测真实文件类型
    
    返回类型: xlsx, xls, csv, docx, doc, png, jpg, gif, bmp, pdf, json, unknown
    """
    if not file_bytes or len(file_bytes) < 4:
        return _guess_type_from_extension(filename)
    
    if file_bytes[:4] == b'\x89PNG':
        return 'png'
    if file_bytes[:3] == b'\xff\xd8\xff':
        return 'jpg'
    if file_bytes[:4] == b'GIF8':
        return 'gif'
    if file_bytes[0] == 0x42 and file_bytes[1] == 0x4D:
        return 'bmp'
    if file_bytes[:4] in (b'II*\x00', b'MM\x00*'):
        return 'tiff'
    if file_bytes[:5] == b'%PDF-':
        return 'pdf'
    # 检查 .xls 文件 (CFBF 格式，Excel 97-2003)
    if len(file_bytes) >= 8 and file_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        return 'xls'
    if file_bytes[:2] == b'PK':
        return _detect_zip_type(file_bytes)
    if _looks_like_csv(file_bytes):
        return 'csv'
    if _looks_like_json(file_bytes):
        return 'json'
    return _guess_type_from_extension(filename)


def _detect_zip_type(file_bytes: bytes) -> str:
    """检测 ZIP 内部文件类型"""
    try:
        import zipfile
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = zf.namelist()
            if any('xl/workbook.xml' in n or 'xl/worksheets/' in n for n in names):
                return 'xlsx'
            if any('word/document.xml' in n for n in names):
                return 'docx'
            if any('ppt/presentation.xml' in n for n in names):
                return 'pptx'
    except:
        pass
    return 'unknown'


def _looks_like_csv(file_bytes: bytes) -> bool:
    """检查是否像 CSV 文件"""
    try:
        text = file_bytes[:2000].decode('utf-8', errors='ignore')
        lines = text.split('\n')[:5]
        if not lines:
            return False
        first = lines[0]
        comma_count = first.count(',')
        semicolon_count = first.count(';')
        tab_count = first.count('\t')
        return comma_count >= 2 or semicolon_count >= 2 or tab_count >= 1
    except:
        return False


def _looks_like_json(file_bytes: bytes) -> bool:
    """检查是否像 JSON 文件"""
    try:
        text = file_bytes[:100].decode('utf-8', errors='ignore').strip()
        return text.startswith('{') or text.startswith('[')
    except:
        return False


def _guess_type_from_extension(filename: str) -> str:
    """根据文件扩展名猜测类型"""
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    mapping = {
        'xlsx': 'xlsx', 'xls': 'xls', 'csv': 'csv', 'tsv': 'csv',
        'docx': 'docx', 'doc': 'doc',
        'png': 'png', 'jpg': 'jpg', 'jpeg': 'jpg', 'gif': 'gif', 'bmp': 'bmp', 'tiff': 'tiff', 'tif': 'tiff',
        'pdf': 'pdf', 'json': 'json'
    }
    return mapping.get(ext, 'unknown')


# ============ Excel/CSV 解析 ============
def parse_excel_or_csv(file_bytes: bytes, filename: str, header_row: int = 0, sheet_name: str = None) -> Dict[str, Any]:
    """
    解析 Excel 或 CSV 文件
    
    Returns:
        {
            'success': True,
            'all_rows': [...],  # 所有行数据
            'potential_headers': [{'row': 0, 'cells': [...]}, ...],  # 可能的表头行
            'detected_header_row': 0,  # 智能检测的表头行
            'file_type': 'xlsx' | 'xls' | 'csv',
            'sheet_names': [...],
            'source': 'excel'
        }
    """
    if not EXCEL_AVAILABLE:
        raise Exception("未安装 openpyxl/pandas 库。请运行: pip install openpyxl pandas")
    
    if not file_bytes or len(file_bytes) == 0:
        raise Exception("上传的文件为空")
    
    file_type = detect_file_type(file_bytes, filename)
    
    if file_type == 'csv':
        return _parse_csv(file_bytes)
    elif file_type in ('xlsx', 'xls'):
        return _parse_xlsx(file_bytes, sheet_name, file_type)
    else:
        raise Exception(
            f"无法解析此文件格式。\n\n"
            f"检测到的文件类型: {file_type}\n"
            f"支持的格式: .xlsx, .xls, .csv\n\n"
            f"请将文件转换为 Excel 格式后重试。"
        )


def _is_potential_header_row(cells: List[str], row_index: int) -> Tuple[bool, float]:
    """
    判断某一行是否可能是表头行
    
    Returns:
        (is_potential, confidence) - 是否可能、置信度
    """
    if not cells:
        return False, 0.0
    
    # 过滤空单元格
    non_empty = [c for c in cells if c and c.strip()]
    if not non_empty:
        return False, 0.0
    
    # 统计
    total_cells = len(cells)
    non_empty_count = len(non_empty)
    empty_ratio = 1 - non_empty_count / total_cells
    
    # 置信度基础分
    confidence = 0.0
    
    # 1. 空单元格比例不宜太高（表头行通常比较完整）
    if empty_ratio < 0.3:
        confidence += 0.3
    elif empty_ratio < 0.5:
        confidence += 0.1
    
    # 2. 非空单元格数量（表头通常有一定数量的列）
    if 2 <= non_empty_count <= 20:
        confidence += 0.3
    elif non_empty_count > 20:
        confidence += 0.2
    
    # 3. 检查是否像字段名（中文、英文、混合）
    field_name_count = 0
    for cell in non_empty:
        cell = cell.strip()
        # 中文字段名
        if re.search(r'[\u4e00-\u9fa5]', cell) and 1 <= len(cell) <= 30:
            field_name_count += 1
        # 英文字段名（无空格、驼峰或下划线）
        elif re.match(r'^[a-zA-Z][a-zA-Z0-9_]*$', cell) and 2 <= len(cell) <= 30:
            field_name_count += 1
        # 短英文单词
        elif re.match(r'^[a-zA-Z]{2,15}$', cell):
            field_name_count += 0.5
    
    if field_name_count >= non_empty_count * 0.6:
        confidence += 0.4
    
    # 4. 检查是否像数据行（全是数字或长文本）
    data_indicators = 0
    for cell in non_empty[:5]:  # 只检查前5个
        cell = cell.strip()
        # 纯数字
        if re.match(r'^[\d,\.]+$', cell):
            data_indicators += 1
        # 长文本（可能是备注等）
        elif len(cell) > 50:
            data_indicators += 1
    
    if data_indicators >= 3:
        confidence -= 0.3
    
    # 5. 检查是否像标题行（很长、包含特定词汇）
    title_keywords = ['标题', '标题', 'title', '报表', '清单', '列表', '目录']
    for cell in non_empty[:2]:
        if any(kw in cell.lower() for kw in title_keywords):
            confidence -= 0.3
            break
    
    return confidence >= 0.4, min(confidence, 1.0)


def _parse_csv(file_bytes: bytes) -> Dict[str, Any]:
    """解析 CSV 文件"""
    encodings = ['utf-8-sig', 'utf-8', 'gbk', 'gb2312', 'gb18030']
    
    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        raise Exception("CSV 文件编码不受支持，请保存为 UTF-8 或 GBK 格式")
    
    lines = text.strip().split('\n')
    if not lines:
        raise Exception("CSV 文件内容为空")
    
    # 检测分隔符
    first_line = lines[0]
    delimiter = ','
    if '\t' in first_line:
        delimiter = '\t'
    elif ';' in first_line:
        delimiter = ';'
    
    # 解析所有行
    all_rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        cells = _split_csv_line(line, delimiter)
        all_rows.append(cells)
    
    if len(all_rows) < 2:
        raise Exception("CSV 文件数据行不足（需要至少1行表头 + 1行数据）")
    
    # 智能检测可能的表头行
    potential_headers, detected_row = _detect_header_rows(all_rows)
    
    return {
        'success': True,
        'all_rows': all_rows,
        'potential_headers': potential_headers,
        'detected_header_row': detected_row,
        'file_type': 'csv',
        'sheet_names': [],
        'source': 'excel'
    }


def _detect_header_rows(all_rows: List[List[str]]) -> Tuple[List[Dict], int]:
    """
    智能检测可能的表头行
    
    Returns:
        (potential_headers, best_row) - 可能的表头行列表、最佳行索引
    """
    if not all_rows or len(all_rows) < 2:
        return [], 0
    
    potential_headers = []
    best_row = 0
    best_confidence = 0.0
    
    # 检查前10行
    max_check_rows = min(10, len(all_rows))
    
    for row_idx in range(max_check_rows):
        cells = all_rows[row_idx]
        is_potential, confidence = _is_potential_header_row(cells, row_idx)
        
        # 生成行的预览文本（最多显示前5个单元格）
        preview = [c.strip() for c in cells[:5] if c.strip()]
        preview_text = ' | '.join(preview) if preview else '(空行)'
        if len(cells) > 5:
            preview_text += f' ... (共{len(cells)}列)'
        
        entry = {
            'row': row_idx,
            'cells': cells,
            'preview': preview_text,
            'is_potential': is_potential,
            'confidence': round(confidence, 2)
        }
        
        potential_headers.append(entry)
        
        if is_potential and confidence > best_confidence:
            best_confidence = confidence
            best_row = row_idx
    
    return potential_headers, best_row


def _split_csv_line(line: str, delimiter: str = ',') -> List[str]:
    """正确分割 CSV 行，处理引号"""
    result = []
    current = []
    in_quote = False
    
    for char in line:
        if char == '"':
            in_quote = not in_quote
        elif char == delimiter and not in_quote:
            result.append(''.join(current).strip())
            current = []
        else:
            current.append(char)
    
    result.append(''.join(current).strip())
    return result


def _parse_xlsx(file_bytes: bytes, sheet_name: str = None, file_type: str = 'xlsx') -> Dict[str, Any]:
    """解析 Excel 文件"""
    errors = []
    
    # 方法1: openpyxl (仅支持 .xlsx)
    if file_type == 'xlsx':
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
            sheet_names = wb.sheetnames
            
            if not sheet_names:
                raise Exception("Excel 文件没有工作表")
            
            # 选择工作表
            if sheet_name and sheet_name in sheet_names:
                ws = wb[sheet_name]
            else:
                ws = wb.active
            
            # 读取数据
            all_rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else '' for cell in row]
                if any(c.strip() for c in cells):
                    all_rows.append(cells)
            
            wb.close()
            
            if not all_rows:
                raise Exception("Excel 文件数据为空")
            
            # 智能检测可能的表头行
            potential_headers, detected_row = _detect_header_rows(all_rows)
            
            return {
                'success': True,
                'all_rows': all_rows,
                'potential_headers': potential_headers,
                'detected_header_row': detected_row,
                'file_type': 'xlsx',
                'sheet_names': sheet_names,
                'current_sheet': sheet_name or wb.sheetnames[0] if hasattr(wb, 'sheetnames') else '',
                'source': 'excel'
            }
        except Exception as e:
            errors.append(f"openpyxl: {str(e)}")
    
    # 方法2: pandas
    try:
        # 根据文件类型选择合适的引擎
        if file_type == 'xls' and XLRD_AVAILABLE:
            # .xls 文件需要用 xlrd 引擎
            df = pd.read_excel(io.BytesIO(file_bytes), header=None, dtype=str, engine='xlrd')
        else:
            # .xlsx 文件用 openpyxl 引擎，或者不指定引擎让 pandas 自动选择
            try:
                df = pd.read_excel(io.BytesIO(file_bytes), header=None, dtype=str, engine='openpyxl')
            except Exception:
                # 如果 openpyxl 失败，尝试不指定引擎
                df = pd.read_excel(io.BytesIO(file_bytes), header=None, dtype=str)
        
        df = df.fillna('')
        
        all_rows = []
        for _, row in df.iterrows():
            cells = [str(v).strip() for v in row.values]
            if any(c.strip() for c in cells):
                all_rows.append(cells)
        
        if not all_rows:
            raise Exception("pandas 读取结果为空")
        
        potential_headers, detected_row = _detect_header_rows(all_rows)
        
        return {
            'success': True,
            'all_rows': all_rows,
            'potential_headers': potential_headers,
            'detected_header_row': detected_row,
            'file_type': file_type,
            'sheet_names': [],
            'source': 'excel'
        }
    except Exception as e:
        errors.append(f"pandas: {str(e)}")
    
    # 方法3: xlrd (仅 .xls)
    try:
        if XLRD_AVAILABLE:
            workbook = xlrd.open_workbook(file_contents=file_bytes)
            sheet = workbook.sheet_by_index(0)
            
            all_rows = []
            for row_idx in range(sheet.nrows):
                cells = [str(sheet.cell(row_idx, col_idx).value) for col_idx in range(sheet.ncols)]
                if any(c.strip() for c in cells):
                    all_rows.append(cells)
            
            if not all_rows:
                raise Exception("xlrd 读取结果为空")
            
            potential_headers, detected_row = _detect_header_rows(all_rows)
            
            return {
                'success': True,
                'all_rows': all_rows,
                'potential_headers': potential_headers,
                'detected_header_row': detected_row,
                'file_type': 'xls',
                'sheet_names': workbook.sheet_names(),
                'source': 'excel'
            }
    except Exception as e:
        errors.append(f"xlrd: {str(e)}")
    
    # 所有方法失败
    error_msg = "\n".join(errors)
    raise Exception(
        f"Excel 文件解析失败。\n\n"
        f"可能原因：\n"
        f"1. 文件已损坏\n"
        f"2. 文件格式不兼容\n\n"
        f"建议：请用 Excel/WPS 打开后另存为新的 .xlsx 文件。\n\n"
        f"详细信息：{error_msg[:200]}"
    )


# ============ Word 解析 ============
def parse_word(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    解析 Word 文档 (.docx)，提取文本和表格
    
    Returns:
        {
            'success': True,
            'potential_headers': [...],  # 提取的标题/关键词作为候选
            'detected_header_row': 0,
            'all_rows': [[...]],  # 候选行列表
            'file_type': 'docx',
            'source': 'word',
            'extracted_text': '...'
        }
    """
    if not WORD_AVAILABLE:
        raise Exception(
            "未安装 python-docx 库。请运行: pip install python-docx\n\n"
            "Word 文档解析需要此依赖。"
        )
    
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        
        # 提取所有表格
        all_rows = []
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text)
                
                if any(cells):  # 只添加非空行
                    all_rows.append(cells)
        
        # 如果没有表格，提取段落
        if not all_rows:
            # 方法1: docx2txt
            try:
                text = docx2txt.process(io.BytesIO(file_bytes))
            except:
                text = None
            
            if not text:
                text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if not text.strip():
                raise Exception("Word 文档内容为空")
            
            # 收集所有非空段落
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            for idx, para_text in enumerate(paragraphs):
                # 标题样式优先
                is_heading = False
                if paragraphs and idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    if para.style and 'Heading' in str(para.style.name):
                        is_heading = True
                
                # 判断是否像表头
                is_potential = False
                confidence = 0.0
                
                # 短行可能是表头
                if 2 <= len(para_text) <= 30:
                    is_potential = True
                    confidence = 0.6
                    if not para_text.endswith(('.', '。', ',', '，', '!', '！', '?', '？')):
                        confidence += 0.2
                
                # 标题样式高置信度
                if is_heading:
                    confidence = 0.9
                
                # 构建单元格（按标点或空格分割）
                cells = re.split(r'[,，\t]', para_text)
                cells = [c.strip() for c in cells if c.strip()]
                
                if not cells:
                    continue
                
                all_rows.append(cells)
                
                # 最多处理100个段落
                if idx >= 99:
                    break
            
            # 使用 jieba 提取关键词作为额外候选
            if JIEBA_AVAILABLE and len(all_rows) < 5:
                try:
                    keywords = jieba.analyse.extract_tags(text, topK=15, withWeight=False)
                    keywords = [k for k in keywords if len(k) >= 2]
                    
                    for i, kw in enumerate(keywords):
                        if kw not in all_rows:
                            entry = {
                                'row': len(all_rows),
                                'cells': [kw],
                                'preview': kw,
                                'is_potential': True,
                                'confidence': 0.5
                            }
                            potential_headers.append(entry)
                            all_rows.append([kw])
                except:
                    pass
        
        if not all_rows:
            raise Exception("Word 文档中未找到可识别的表格数据")
        
        # 生成候选表头
        potential_headers = []
        detected_row = 0
        
        for i, row in enumerate(all_rows[:10]):
            if len(row) >= 2 and any(c.strip() for c in row):
                entry = {
                    'row': i,
                    'cells': row,
                    'preview': ' | '.join(row[:5]),
                    'is_potential': True,
                    'confidence': 0.8 if i == 0 else 0.5
                }
                potential_headers.append(entry)
        
        return {
            'success': True,
            'potential_headers': potential_headers,
            'detected_header_row': detected_row,
            'all_rows': all_rows[:100],
            'file_type': 'docx',
            'source': 'word',
            'extracted_text': f"从 {table_count} 个表格中提取" if table_count > 0 else "",
            'sheet_names': []
        }
        
    except Exception as e:
        error_msg = str(e)
        if 'zipfile' in error_msg.lower() or 'corrupt' in error_msg.lower():
            raise Exception(
                f"Word 文件可能已损坏。\n\n"
                f"建议：\n"
                f"1. 用 Word/WPS 重新打开文件\n"
                f"2. 另存为新的 .docx 文件\n"
                f"3. 重新上传"
            )
        raise Exception(f"Word 文档解析失败：{error_msg}")


# ============ PDF 解析 ============
def parse_pdf(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    解析 PDF 文件，提取文本和表格数据
    
    Returns:
        {
            'success': True,
            'potential_headers': [...],
            'detected_header_row': 0,
            'all_rows': [...],
            'file_type': 'pdf',
            'source': 'pdf',
            'extracted_text': '...'
        }
    """
    try:
        import pdfplumber
    except ImportError:
        raise Exception(
            "PDF 解析功能需要安装 pdfplumber 库。\n\n"
            "请在终端运行：pip install pdfplumber"
        )
    
    try:
        pdf_file = io.BytesIO(file_bytes)
        all_rows = []
        potential_headers = []
        full_text = ""
        
        with pdfplumber.open(pdf_file) as pdf:
            page_count = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # 尝试提取表格
                tables = page.extract_tables()
                
                if tables:
                    for table_idx, table in enumerate(tables):
                        if not table:
                            continue
                        
                        # 处理表格数据
                        for row_idx, row in enumerate(table):
                            if not row:
                                continue
                            
                            # 清理单元格数据
                            cells = []
                            for cell in row:
                                if cell is None:
                                    cells.append('')
                                else:
                                    # 清理单元格文本
                                    cell_text = str(cell).strip()
                                    # 移除多余的空白字符
                                    cell_text = re.sub(r'\s+', ' ', cell_text)
                                    cells.append(cell_text)
                            
                            if any(cells):  # 只添加非空行
                                all_rows.append(cells)
                
                # 同时提取纯文本作为备选
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n\n"
            
            if not all_rows:
                # 如果没有提取到表格，尝试从文本解析表格结构
                all_rows = _parse_text_as_table(full_text)
        
        if not all_rows:
            raise Exception(
                "PDF 文件中未找到可识别的表格数据。\n\n"
                "建议：\n"
                "1. 如果 PDF 是扫描件，请将其截图后上传图片\n"
                "2. 将 PDF 中的表格复制到 Excel，然后导入 Excel\n"
                "3. 使用专业的 PDF 表格提取工具"
            )
        
        # 生成候选表头
        detected_row = 0
        for i, row in enumerate(all_rows[:10]):
            if len(row) >= 2 and any(c.strip() for c in row):
                entry = {
                    'row': i,
                    'cells': row,
                    'preview': ' | '.join(row[:5]),
                    'is_potential': True,
                    'confidence': 0.8 if i == 0 else 0.5
                }
                potential_headers.append(entry)
        
        return {
            'success': True,
            'potential_headers': potential_headers,
            'detected_header_row': detected_row,
            'all_rows': all_rows[:100],
            'file_type': 'pdf',
            'source': 'pdf',
            'extracted_text': full_text[:5000],
            'sheet_names': []
        }
        
    except Exception as e:
        error_msg = str(e)
        if 'password' in error_msg.lower():
            raise Exception("PDF 文件已加密，需要密码解锁")
        if 'corrupt' in error_msg.lower() or 'invalid' in error_msg.lower():
            raise Exception(
                f"PDF 文件可能已损坏。\n\n"
                f"建议：\n"
                f"1. 用 PDF 阅读器重新打开文件\n"
                f"2. 另存为新的 PDF 文件\n"
                f"3. 或者将表格复制到 Excel 后导入"
            )
        raise Exception(f"PDF 解析失败：{error_msg}")


def _parse_text_as_table(text: str) -> List[List[str]]:
    """
    尝试从纯文本中解析表格结构
    """
    rows = []
    
    # 按行分割
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 尝试多种分隔符
        cells = None
        
        # 尝试 Tab 分隔
        if '\t' in line:
            cells = [c.strip() for c in line.split('\t') if c.strip()]
        
        # 尝试多个空格分隔
        if not cells and '  ' in line:
            cells = [c.strip() for c in re.split(r'\s{2,}', line) if c.strip()]
        
        # 尝试逗号分隔
        if not cells and ',' in line:
            cells = [c.strip() for c in line.split(',') if c.strip()]
        
        if cells and len(cells) >= 2:
            rows.append(cells)
    
    return rows


# ============ 图片 OCR 解析 ============
def parse_image(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    使用 OCR 从图片中提取文字和表格
    
    Returns:
        {
            'success': True,
            'potential_headers': [...],
            'detected_header_row': 0,
            'all_rows': [...],
            'file_type': 'png' | 'jpg' | ...,
            'source': 'ocr',
            'extracted_text': '...'
        }
    """
    if not TESSERACT_AVAILABLE:
        raise Exception(
            "OCR 功能未就绪。\n\n"
            "请在终端运行以下命令安装：\n"
            "pip install pytesseract Pillow opencv-python\n\n"
            "同时需要安装 Tesseract OCR 引擎：\n"
            "下载地址: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    
    tesseract_cmd = _get_tesseract_path()
    
    if not tesseract_cmd or not os.path.exists(tesseract_cmd):
        raise Exception(
            "Tesseract OCR 引擎未配置。\n\n"
            "请在【系统配置】页面设置 Tesseract 路径。\n"
            "下载地址: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    
    try:
        import pytesseract
        from PIL import Image
        
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        
        # 打开图片
        try:
            pil_img = Image.open(io.BytesIO(file_bytes))
            if pil_img.mode != 'RGB':
                pil_img = pil_img.convert('RGB')
        except Exception as e:
            raise Exception(f"无法读取图片文件：{str(e)}")
        
        # 图像预处理
        processed_img = pil_img
        if CV2_AVAILABLE:
            try:
                import numpy as np
                img_array = np.array(pil_img)
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                processed_img = Image.fromarray(binary)
            except:
                pass
        
        # 检查中文语言包
        lang = 'eng'
        try:
            test_img = Image.new('RGB', (50, 20), color='white')
            pytesseract.image_to_string(test_img, lang='chi_sim+eng')
            lang = 'chi_sim+eng'
        except:
            pass
        
        # OCR 识别
        custom_config = r'--oem 3 --psm 6'
        try:
            text = pytesseract.image_to_string(processed_img, lang=lang, config=custom_config)
        except Exception as e:
            err_str = str(e).lower()
            if 'language' in err_str or 'data' in err_str or 'chi_sim' in err_str:
                text = pytesseract.image_to_string(processed_img, lang='eng', config=custom_config)
                lang = 'eng'
            else:
                raise
        
        if not text.strip():
            raise Exception("未能从图片中识别到文字，请确保图片清晰且包含文字内容")
        
        # 解析表格结构
        headers, rows = _parse_ocr_table(text)
        
        # 构建返回数据
        potential_headers = []
        all_rows = []
        
        if headers:
            # 第一行作为表头
            entry = {
                'row': 0,
                'cells': headers,
                'preview': ' | '.join(headers[:5]),
                'is_potential': True,
                'confidence': 0.8
            }
            potential_headers.append(entry)
            all_rows.append(headers)
        
        # 其他数据行
        for i, row in enumerate(rows[:20]):
            if row and row != headers:
                entry = {
                    'row': len(all_rows),
                    'cells': row,
                    'preview': ' | '.join(row[:5]),
                    'is_potential': False,
                    'confidence': 0.3
                }
                potential_headers.append(entry)
                all_rows.append(row)
        
        # 如果没识别到表格，提取关键词作为表头
        if not headers:
            keywords = []
            if JIEBA_AVAILABLE:
                try:
                    keywords = jieba.analyse.extract_tags(text, topK=10, withWeight=False)
                except:
                    pass
            
            if not keywords:
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if 2 <= len(line) <= 15:
                        keywords.append(line)
            
            headers = keywords[:10] if keywords else ['字段1', '字段2', '字段3']
            
            entry = {
                'row': 0,
                'cells': headers,
                'preview': ' | '.join(headers[:5]),
                'is_potential': True,
                'confidence': 0.5
            }
            potential_headers.append(entry)
            all_rows.append(headers)
        
        return {
            'success': True,
            'potential_headers': potential_headers,
            'detected_header_row': 0,
            'all_rows': all_rows,
            'file_type': detect_file_type(file_bytes, filename),
            'source': 'ocr',
            'extracted_text': text[:3000]
        }
        
    except Exception as e:
        error_msg = str(e)
        if 'language' in error_msg.lower() or 'chi_sim' in error_msg.lower():
            raise Exception(
                "Tesseract 缺少中文语言包。\n\n"
                "请重新安装 Tesseract，安装时勾选 Chinese Simplified (chi_sim)。\n"
                "下载地址: https://github.com/UB-Mannheim/tesseract/wiki"
            )
        if 'not found' in error_msg.lower() or '不存在' in error_msg:
            raise Exception(
                f"Tesseract 路径不存在：{tesseract_cmd}\n\n"
                "请检查【系统配置】中的 Tesseract 路径是否正确。"
            )
        raise Exception(f"图片 OCR 识别失败：{error_msg}")


def _parse_ocr_table(text: str) -> Tuple[List[str], List[List[str]]]:
    """解析 OCR 识别的表格文本"""
    headers = []
    rows = []
    
    if not text:
        return headers, rows
    
    lines = text.strip().split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # 尝试多种分隔符
        cells = re.split(r'[\t│┌┐└┘├┤┼─=]{2,}', line)
        if len(cells) < 2:
            cells = re.split(r'\s{2,}', line)
        if len(cells) < 2:
            cells = re.split(r'\t', line)
        
        cells = [c.strip() for c in cells if c.strip()]
        
        if len(cells) >= 2:
            if i == 0 or not headers:
                headers = cells
            else:
                rows.append(cells)
    
    if headers:
        headers = [h for h in headers if h][:20]
    
    if rows:
        rows = [r for r in rows if len(r) >= 2 and any(c for c in r)][:50]
    
    return headers, rows


def _get_tesseract_path() -> str:
    """获取 Tesseract 路径"""
    # 1. 数据库配置
    try:
        from app.core.config import settings
        db_path = settings.DATABASE_URL
        if db_path and 'sqlite' in db_path:
            try:
                import sqlite3
                if ':///' in db_path:
                    db_file = db_path.split(':///')[-1]
                else:
                    db_file = db_path.split('://')[1]
                
                if os.path.exists(db_file):
                    conn = sqlite3.connect(db_file)
                    cursor = conn.cursor()
                    cursor.execute("SELECT value FROM system_configs WHERE key = 'ocr_tesseract_path' AND organization_id IS NULL")
                    row = cursor.fetchone()
                    conn.close()
                    if row and row[0]:
                        return row[0]
            except:
                pass
    except:
        pass
    
    # 2. 环境变量
    path = os.environ.get('TESSERACT_CMD') or os.environ.get('TESSERACT_PATH', '')
    if path and os.path.exists(path):
        return path
    
    # 3. settings
    try:
        from app.core.config import settings
        path = getattr(settings, 'TESSERACT_PATH', '') or ''
        if path and os.path.exists(path):
            return path
    except:
        pass
    
    # 4. 常见路径
    common_paths = [
        'C:\\Program Files\\Tesseract-OCR\\tesseract.exe',
        'C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe',
        'C:\\Tesseract-OCR\\tesseract.exe',
        '/usr/bin/tesseract',
        '/usr/local/bin/tesseract',
    ]
    for p in common_paths:
        if os.path.exists(p):
            return p
    
    return ''


# ============ 统一解析入口 ============
def parse_file(file_bytes: bytes, filename: str, header_row: int = 0, sheet_name: str = None) -> Dict[str, Any]:
    """
    统一文件解析入口 - 自动检测文件类型并调用对应解析器
    
    注意：此函数返回原始数据和候选表头行，具体使用哪一行作为表头由前端决定
    """
    if not file_bytes or len(file_bytes) == 0:
        raise Exception("上传的文件为空，请选择文件后重试")
    
    file_type = detect_file_type(file_bytes, filename)
    
    if file_type in ('xlsx', 'xls', 'csv'):
        return parse_excel_or_csv(file_bytes, filename, header_row, sheet_name)
    
    elif file_type == 'docx':
        return parse_word(file_bytes, filename)
    
    elif file_type in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff'):
        return parse_image(file_bytes, filename)
    
    elif file_type == 'json':
        raise Exception("JSON 文件请使用【JSON导入】功能")
    
    elif file_type == 'pdf':
        return parse_pdf(file_bytes, filename)
    
    elif file_type == 'doc':
        raise Exception(
            ".doc 格式暂不支持直接导入。\n\n"
            "请将文件另存为 .docx 格式后重试：\n"
            "用 Word/WPS 打开文件 → 文件 → 另存为 → 选择 .docx 格式"
        )
    
    else:
        raise Exception(
            f"不支持的文件格式。\n\n"
            f"检测到的类型: {file_type}\n\n"
            f"支持的格式：\n"
            f"• Excel: .xlsx, .xls, .csv\n"
            f"• Word: .docx\n"
            f"• 图片: .png, .jpg, .jpeg, .bmp\n\n"
            f"请转换文件格式后重试。"
        )


def apply_header_row(result: Dict[str, Any], header_row: int) -> Dict[str, Any]:
    """
    根据选定的表头行构建headers和rows
    
    Args:
        result: parse_file 返回的原始结果
        header_row: 用户选择的表头行索引
    
    Returns:
        {
            'headers': [...],
            'rows': [...],
            'all_rows': [...]
        }
    """
    all_rows = result.get('all_rows', [])
    
    if not all_rows:
        return {'headers': [], 'rows': [], 'all_rows': []}
    
    # 确保索引有效
    header_row = max(0, min(header_row, len(all_rows) - 1))
    
    headers = all_rows[header_row]
    rows = all_rows[header_row + 1:]
    
    return {
        'headers': headers,
        'rows': rows,
        'all_rows': all_rows
    }


# ============ 字段生成 ============
FIELD_TYPE_KEYWORDS = {
    'text': ['姓名', '名称', '公司', '地址', '备注', '说明', '描述', '标题', '职位', '部门', '联系人', '品牌', '产品', '编码', '编号'],
    'number': ['数量', '金额', '价格', '单价', '总价', '折扣', '税率', '面积', '重量'],
    'phone': ['电话', '手机', '固话', '传真'],
    'email': ['邮箱', '邮件'],
    'date': ['日期', '时间', '生日', '到期', '有效期'],
    'select': ['类型', '分类', '状态', '等级', '方式', '渠道'],
    'money': ['工资', '薪酬', '预算', '成本', '利润', '收入', '支出'],
}

FIELD_NAME_MAP = {
    '姓名': 'name', '名称': 'name', '公司': 'company', '地址': 'address',
    '电话': 'phone', '手机': 'mobile', '邮箱': 'email', '日期': 'date',
    '数量': 'quantity', '金额': 'amount', '价格': 'price', '单价': 'unit_price',
    '类型': 'type', '分类': 'category', '状态': 'status', '等级': 'level',
    '备注': 'remark', '说明': 'description', '编码': 'code', '编号': 'code',
    '职位': 'position', '部门': 'department', '联系人': 'contact',
    '产品': 'product', '规格': 'spec', '品牌': 'brand', '单位': 'unit',
}


def infer_field_type(header: str, values: List[str] = None) -> str:
    """根据表头名称推断字段类型"""
    header_lower = header.lower()
    header_clean = header.strip()
    
    for ftype, keywords in FIELD_TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in header_clean or kw in header_lower:
                return ftype
    
    if values:
        non_empty = [str(v).strip() for v in values if v and str(v).strip()]
        if not non_empty:
            return 'text'
        
        numeric_count = sum(1 for v in non_empty[:20] if _is_numeric(v))
        if numeric_count > len(non_empty) * 0.7:
            return 'number'
        
        date_count = sum(1 for v in non_empty[:20] if _is_date(v))
        if date_count > len(non_empty) * 0.7:
            return 'date'
        
        if sum(1 for v in non_empty if '@' in v and '.' in v) > len(non_empty) * 0.5:
            return 'email'
        
        if sum(1 for v in non_empty if _is_phone(v)) > len(non_empty) * 0.5:
            return 'phone'
        
        unique_vals = set(non_empty[:50])
        if 1 < len(unique_vals) <= 10:
            return 'select'
    
    return 'text'


def _is_numeric(s: str) -> bool:
    s = str(s).replace(',', '').replace('¥', '').replace('$', '').replace('%', '').strip()
    try:
        float(s)
        return True
    except ValueError:
        return False


def _is_date(s: str) -> bool:
    s = str(s).strip()
    patterns = [r'^\d{4}-\d{1,2}-\d{1,2}', r'^\d{4}/\d{1,2}/\d{1,2}', 
                r'^\d{1,2}-\d{1,2}-\d{4}', r'^\d{1,2}/\d{1,2}/\d{4}']
    return any(re.match(p, s) for p in patterns)


def _is_phone(s: str) -> bool:
    s = str(s).replace('-', '').replace(' ', '').replace('(', '').replace(')', '')
    return bool(re.match(r'^1[3-9]\d{9}$', s)) or bool(re.match(r'^\d{7,12}$', s))


def chinese_to_pinyin(text: str) -> str:
    """中文转拼音字段名"""
    if not text:
        return 'field'
    
    if text in FIELD_NAME_MAP:
        return FIELD_NAME_MAP[text]
    
    if JIEBA_AVAILABLE:
        words = jieba.cut(text)
        for w in words:
            w = w.strip()
            if w in FIELD_NAME_MAP:
                return FIELD_NAME_MAP[w]
    
    result = re.sub(r'[^a-z0-9\u4e00-\u9fa5]', '', text)
    if not result:
        return 'field'
    
    pinyin = ''
    for char in result[:6]:
        for key, val in FIELD_NAME_MAP.items():
            if key.startswith(char):
                pinyin += val
                break
        else:
            pinyin += char
    
    pinyin = re.sub(r'[^a-z0-9_]', '_', pinyin.lower())
    pinyin = re.sub(r'_+', '_', pinyin).strip('_')
    return pinyin or 'field'


def generate_fields(headers: List[str], rows: List[List[str]] = None) -> List[Dict[str, Any]]:
    """根据表头和示例数据生成字段定义"""
    fields = []
    used_names = set()
    
    for i, header in enumerate(headers):
        header = header.strip()
        if not header:
            continue
        
        base_name = chinese_to_pinyin(header)
        field_name = base_name
        counter = 1
        while field_name in used_names:
            field_name = f"{base_name}_{counter}"
            counter += 1
        used_names.add(field_name)
        
        col_values = []
        if rows:
            for row in rows:
                if i < len(row):
                    col_values.append(row[i])
        
        field_type = infer_field_type(header, col_values)
        
        options = []
        if field_type == 'select' and col_values:
            unique_vals = list(set(v.strip() for v in col_values if v.strip()))[:20]
            options = unique_vals
        
        field = {
            'name': field_name,
            'label': header,
            'type': field_type,
            'placeholder': f'请输入{header}' if field_type == 'text' else '',
            'required': False,
            'width': '100%',
            'options': options,
            'optionsText': '，'.join(options),
            'maxLength': 255,
        }
        fields.append(field)
    
    return fields


# ============ 状态检查 ============
def get_dependencies_status() -> Dict[str, Any]:
    """获取所有依赖组件的状态"""
    tesseract_status = {
        'available': TESSERACT_AVAILABLE,
        'engine_configured': False,
        'chi_sim_installed': False,
        'message': ''
    }
    
    if TESSERACT_AVAILABLE:
        try:
            tesseract_cmd = _get_tesseract_path()
            if tesseract_cmd and os.path.exists(tesseract_cmd):
                import pytesseract
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                
                try:
                    version = pytesseract.get_tesseract_version()
                    tesseract_status['engine_configured'] = True
                    
                    try:
                        from PIL import Image
                        test_img = Image.new('RGB', (50, 20), color='white')
                        pytesseract.image_to_string(test_img, lang='chi_sim+eng')
                        tesseract_status['chi_sim_installed'] = True
                        tesseract_status['message'] = f'已就绪 (v{version})'
                    except:
                        tesseract_status['message'] = '引擎已就绪，缺少中文语言包'
                except Exception as e:
                    tesseract_status['message'] = f'引擎错误: {str(e)[:50]}'
            else:
                tesseract_status['message'] = '请配置 Tesseract 路径'
        except Exception as e:
            tesseract_status['message'] = f'错误: {str(e)[:50]}'
    
    ocr_ready = tesseract_status['chi_sim_installed']
    
    return {
        'excel': {
            'available': EXCEL_AVAILABLE,
            'message': '已就绪' if EXCEL_AVAILABLE else '未安装'
        },
        'word': {
            'available': WORD_AVAILABLE,
            'message': '已就绪' if WORD_AVAILABLE else '未安装'
        },
        'jieba': {
            'available': JIEBA_AVAILABLE,
            'message': '已就绪' if JIEBA_AVAILABLE else '未安装'
        },
        'ocr': tesseract_status,
        'ocr_ready': ocr_ready,
        'cv2': {
            'available': CV2_AVAILABLE,
            'message': '已就绪' if CV2_AVAILABLE else '未安装'
        }
    }


# ============ 预览构建 ============
def build_preview(headers: List[str], rows: List[List[str]], fields: List[Dict], max_rows: int = 10) -> Dict[str, Any]:
    """构建预览数据"""
    preview_rows = rows[:max_rows]
    preview_data = []
    
    for row in preview_rows:
        row_data = {}
        for j, field in enumerate(fields):
            if j < len(row):
                row_data[field['name']] = row[j]
            else:
                row_data[field['name']] = ''
        preview_data.append(row_data)
    
    return {
        'headers': headers,
        'columns': [f['label'] for f in fields],
        'field_types': [f['type'] for f in fields],
        'rows': preview_data,
        'total_count': len(rows)
    }
